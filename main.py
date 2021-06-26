import traceback
import os
import sys
from PIL import Image, ExifTags
from docx import Document
from docx.shared import Inches
from tkinter import Tk, filedialog

try:
    root = Tk()  # pointing root to Tk() to use it as Tk() in program.
    root.withdraw()  # Hides small tkinter window.

    root.attributes('-topmost', True)  # Opened windows will be active. above all windows despite of selection.

    path = filedialog.askdirectory()  # Returns opened path as str

    desiredExifs = ['GPSInfo', 'Make', 'Model', 'DateTime', 'XResolution', 'YResolution', 'DateTimeOriginal', 'ExifImageWidth', 'ExifImageHeight']

    def renameWithDate(exif, file):
        date = exif.get("DateTimeOriginal")
        date = date[0:date.find(' ')]
        dateNew = date.replace(':', '-')
        newFileName = f'{file[0:file.find(".")]} {dateNew}{file[file.find("."):]}'
        if dateNew in file:
            return file
        else:
            return newFileName

    for file in os.listdir(path):
        print(file)
        if (file != ".DS_Store" and '.docx' not in file) and ('.JPG' in file or '.jpeg' in file or '.jpg' in file or '.JPEG' in file):  # may want to add to check if they are jpeg, JPG, PNG, etc
            img = Image.open(f'{path}/{file}')
            exif = {ExifTags.TAGS[k]: v for k, v in img._getexif().items() if k in ExifTags.TAGS}
            updatedExif = {}
            for i in range(len(desiredExifs)):
                if exif.get(desiredExifs[i]) != None:
                    updatedExif[desiredExifs[i]] = exif.get(desiredExifs[i])
            print(updatedExif)
            '''prints the dictionary
            for k, v in exif.items():
                print(f'{k}: {v}')
            print()
            '''
            document = Document()
            document.add_picture(f'{path}/{file}', width=Inches(6.00))
            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Key'
            hdr_cells[1].text = 'Value'
            for k, v in updatedExif.items():
                row_cells = table.add_row().cells
                row_cells[0].text = f'{k}'
                row_cells[1].text = f'{v}'

            if exif.get("DateTimeOriginal") is None:
                print(f'{file} has no date metadata.\n')
            else:
                action = renameWithDate(exif, file)
                img.close()
                os.rename(f'{path}/{file}', f'{path}/{action}')
            document.save(f'{path}/Metadata {action[0:action.find(".")]}.docx')

except Exception:
    print(traceback.format_exc())
    input("Press return to exit")
